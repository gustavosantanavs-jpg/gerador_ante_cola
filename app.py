import streamlit as st
from docx import Document
import io
import random
import re
import base64

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Sistema Anti-Cola Pro", page_icon="🎓")

def adicionar_fundo_de_tela(arquivo_imagem):
    try:
        with open(arquivo_imagem, "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode()
        st.markdown(
        f'''
        <style>
        .stApp {{
            background-image: url(data:image/png;base64,{encoded_string});
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
        }}
        </style>
        ''', unsafe_allow_html=True)
    except Exception:
        pass 

adicionar_fundo_de_tela("logo.png") 

# --- SISTEMA DE LOGIN ---
if 'logado' not in st.session_state:
    st.session_state['logado'] = False

if not st.session_state['logado']:
    st.markdown('''
        <style>
        .block-container {
            background-color: rgba(255, 255, 255, 0.95);
            padding: 2.5rem;
            border-radius: 15px;
            margin-top: 50vh; 
            max-width: 450px; 
            box-shadow: 0 10px 25px rgba(0, 0, 0, 0.2);
        }
        </style>
    ''', unsafe_allow_html=True)

    st.markdown("<h2 style='text-align: center; color: #1E3A8A;'>Acesso Restrito</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Insira as suas credenciais para aceder ao gerador.</p>", unsafe_allow_html=True)
    
    usuario = st.text_input("Utilizador")
    senha = st.text_input("Palavra-passe", type="password")
    st.write("")
    
    if st.button("Entrar", use_container_width=True):
        if usuario == "milena" and senha == "unimam2026":
            st.session_state['logado'] = True
            st.rerun() 
        else:
            st.error("Utilizador ou palavra-passe incorretos!")
    st.stop()
    # =====================================================================
# --- A PARTIR DAQUI SÓ APARECE SE O LOGIN ESTIVER CORRETO ---
# =====================================================================

st.markdown('''
    <style>
    .block-container {
        background-color: rgba(255, 255, 255, 0.95);
        padding: 3rem;
        border-radius: 15px;
        margin-top: 5vh; 
        max-width: 800px; 
        box-shadow: 0 10px 25px rgba(0, 0, 0, 0.2);
    }
    </style>
''', unsafe_allow_html=True)

st.title("📚 Sistema Anti-Cola Pro - Profa. Milena")
st.write("Faça o upload da prova original em Word (.docx). O sistema irá embaralhar as questões e alternativas.")

# --- CIRURGIA PROFUNDA NO XML (Resolve o fantasma das letras travadas) ---
def atualizar_paragrafo(paragrafo_xml, padrao, novo_texto, aplicar_negrito=False):
    text_elements = paragrafo_xml.findall('.//w:t', namespaces=paragrafo_xml.nsmap)
    texto_completo = ""
    for t in text_elements:
        if t.text:
            texto_completo += t.text
            
    match = padrao.match(texto_completo) 
    if not match: return
    
    tamanho_padrao = match.end()
    texto_acumulado = ""
    runs_modificadas = False
    ultimo_t_alterado = None
    
    for t in text_elements:
        texto_original = t.text
        if not texto_original: 
            continue
            
        if not runs_modificadas:
            texto_acumulado += texto_original
            ultimo_t_alterado = t
            
            if len(texto_acumulado) <= tamanho_padrao:
                t.text = "" 
            else:
                sobra = texto_acumulado[tamanho_padrao:]
                t.text = novo_texto + sobra
                runs_modificadas = True
                
    if not runs_modificadas and ultimo_t_alterado is not None:
        ultimo_t_alterado.text = novo_texto

    if aplicar_negrito and ultimo_t_alterado is not None:
        try:
            r = ultimo_t_alterado.getparent()
            if r is not None and r.tag.endswith('r'):
                rPr = r.find('.//w:rPr', namespaces=r.nsmap)
                if rPr is None:
                    rPr = r.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                    r.insert(0, rPr)
                b = rPr.find('.//w:b', namespaces=r.nsmap)
                if b is None:
                    b = r.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                    rPr.append(b)
        except Exception:
            pass

def processar_prova_com_imagens(doc_original, gerar_gabarito):
    padrao_questao = re.compile(r'^\s*(Questão\s*)?\d+[\.\-\:]?\s*', re.IGNORECASE)
    # Radar ampliado para ir até a letra 'j' caso haja muitas alternativas
    padrao_alternativa = re.compile(r'^\s*[a-j][\)\.\-]\s*', re.IGNORECASE)

    body = doc_original.element.body
    questoes = []
    questao_atual = None
    alternativa_atual = None
    cabecalho = []
    
    for element in list(body):
        if element.tag.endswith('sectPr'):
            continue
        texto = ""
        is_paragraph = element.tag.endswith('p')
        if is_paragraph:
            for t in element.findall('.//w:t', namespaces=element.nsmap):
                if t.text:
                    texto += t.text

        if is_paragraph and padrao_questao.match(texto):
            questao_atual = {'enunciado': [element], 'alternativas': []}
            questoes.append(questao_atual)
            alternativa_atual = None
        elif is_paragraph and padrao_alternativa.match(texto) and questao_atual is not None:
            is_correct = (len(questao_atual['alternativas']) == 0)
            alternativa_atual = {'blocos': [element], 'correta': is_correct}
            questao_atual['alternativas'].append(alternativa_atual)
        else:
            if alternativa_atual is not None:
                alternativa_atual['blocos'].append(element)
            elif questao_atual is not None:
                questao_atual['enunciado'].append(element)
            else:
                cabecalho.append(element)

    random.shuffle(questoes)
    for q in questoes:
        random.shuffle(q['alternativas'])

    sectPr = None
    for child in list(body):
        if child.tag.endswith('sectPr'):
            sectPr = child
        body.remove(child)

    for el in cabecalho:
        body.append(el)

    gabarito_final = {}
    letras_maiusculas = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J']
    letras_formatadas = ['a) ', 'b) ', 'c) ', 'd) ', 'e) ', 'f) ', 'g) ', 'h) ', 'i) ', 'j) ']

    contador_questao = 1
    for q in questoes:
        p_xml = q['enunciado'][0]
        body.append(p_xml)
        atualizar_paragrafo(p_xml, padrao_questao, f"Questão {contador_questao}: ", aplicar_negrito=True)
        for el in q['enunciado'][1:]:
            body.append(el)
        for idx_alt, alt in enumerate(q['alternativas']):
            if alt['correta']:
                gabarito_final[contador_questao] = letras_maiusculas[idx_alt]
            p_xml_alt = alt['blocos'][0]
            body.append(p_xml_alt)
            # Agora injeta a letra nova corretamente ignorando bloqueios
            atualizar_paragrafo(p_xml_alt, padrao_alternativa, letras_formatadas[idx_alt], aplicar_negrito=False)
            for el in alt['blocos'][1:]:
                body.append(el)
        contador_questao += 1

    if gerar_gabarito:
        doc_original.add_page_break() 
        p_titulo = doc_original.add_paragraph()
        run_titulo = p_titulo.add_run("--- GABARITO ---")
        run_titulo.bold = True
        for q_num in range(1, contador_questao):
            if q_num in gabarito_final:
                doc_original.add_paragraph(f"Questão {q_num}: {gabarito_final[q_num]}")

    if sectPr is not None:
        body.append(sectPr)

    return doc_original

arquivo_prova = st.file_uploader("Selecione o ficheiro da prova (.docx)", type=["docx"])

st.info("💡 **Atenção:** Só marque a opção de Gabarito abaixo se montou a prova original colocando todas as respostas corretas na letra 'a)'.")
gerar_gabarito = st.checkbox("Gerar Gabarito Automático no final da prova", value=False)

qtd_versoes = st.number_input("Quantas versões diferentes deseja gerar?", min_value=1, max_value=10, value=2)

if arquivo_prova is not None:
    if st.button("Embaralhar e Gerar Provas"):
        with st.spinner("A embaralhar tudo e a processar os ficheiros..."):
            try:
                for i in range(int(qtd_versoes)):
                    doc_base = Document(arquivo_prova)
                    novo_doc = processar_prova_com_imagens(doc_base, gerar_gabarito)
                    buffer = io.BytesIO()
                    novo_doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label=f"⬇️ Descarregar Prova - Versão {i+1}",
                        data=buffer,
                        file_name=f"Prova_AntiCola_Versao_{i+1}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_{i}"
                    )
                st.success("✨ Sucesso! Provas geradas com formatação original.")
            except Exception as e:
                st.error(f"Ocorreu um erro ao processar o ficheiro. Erro técnico: {e}")
